from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

OUT_DOCX = r"C:\Users\moham\source\repos\Coworkspace\foodie-audit\Foodie_Test_Report.docx"

BLUE = "1F4E79"
DARK_BLUE = RGBColor(31, 78, 121)
DARK_GRAY = RGBColor(64, 64, 64)
LIGHT_GRAY = "F2F4F7"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None, align=None, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_field(paragraph, instr):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instr
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_section_heading(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    shade_paragraph(p, BLUE)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(255, 255, 255)


def add_subheading(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = DARK_GRAY


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    r.bold = bold
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        r.font.name = "Calibri"
        r.font.size = Pt(11)


def add_table(doc, headers, rows, widths=None, severity_cols=None, center_cols=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.allow_autofit = True
    hdr = table.rows[0]
    hdr._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for i, h in enumerate(headers):
        set_cell_shading(hdr.cells[i], BLUE)
        set_cell_margins(hdr.cells[i])
        set_cell_text(hdr.cells[i], h, bold=True, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    sev_map = {
        "CRITICAL": ("8B0000", "FFFFFF"),
        "HIGH": ("F4B183", "000000"),
        "MEDIUM": ("FFF2CC", "000000"),
        "LOW": ("C6E0B4", "000000"),
        "PASS": ("D9EAF7", "000000"),
        "FAIL": ("F4CCCC", "000000"),
        "INFO": ("E7E6E6", "000000"),
        "WARNING": ("FFF2CC", "000000"),
        "PARTIAL": ("FFF2CC", "000000"),
    }
    severity_cols = severity_cols or []
    center_cols = center_cols or []
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for c_idx, value in enumerate(row):
            if r_idx % 2 == 1:
                set_cell_shading(cells[c_idx], "FAFAFA")
            text = str(value)
            align = WD_ALIGN_PARAGRAPH.CENTER if c_idx in center_cols else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[c_idx], text, align=align, size=8.5)
            set_cell_margins(cells[c_idx])
            upper = text.upper()
            if c_idx in severity_cols and upper in sev_map:
                fill, font = sev_map[upper]
                set_cell_shading(cells[c_idx], fill)
                set_cell_text(cells[c_idx], text, bold=True, color=font, align=WD_ALIGN_PARAGRAPH.CENTER, size=8.5)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width
    doc.add_paragraph()
    return table


def add_score_chart(doc, scores):
    add_subheading(doc, "Readiness Score Summary")
    table = doc.add_table(rows=len(scores), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, (label, score) in enumerate(scores):
        cells = table.rows[idx].cells
        set_cell_text(cells[0], label, bold=True, size=9)
        set_cell_text(cells[1], f"{score}/100", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
        bar = "█" * max(1, score // 5) + "░" * (20 - max(1, score // 5))
        set_cell_text(cells[2], bar, color="1F4E79", size=9)
        for cell in cells:
            set_cell_margins(cell)


def setup_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.bold = True
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.color.rgb = DARK_BLUE
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 2"].font.color.rgb = DARK_GRAY
    if "Small Table Text" not in styles:
        st = styles.add_style("Small Table Text", WD_STYLE_TYPE.PARAGRAPH)
        st.font.name = "Calibri"
        st.font.size = Pt(8.5)


def configure_section(section, include_header_footer=True):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    if include_header_footer:
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        header_p = section.header.paragraphs[0]
        header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header_p.text = ""
        r = header_p.add_run("Foodie Platform QA Report")
        r.font.name = "Calibri"
        r.font.size = Pt(9)
        r.font.color.rgb = DARK_GRAY
        footer_p = section.footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer_p.add_run("Page ")
        add_field(footer_p, "PAGE")
        footer_p.add_run(" of ")
        add_field(footer_p, "SECTIONPAGES")


def add_cover(doc):
    sec = doc.sections[0]
    configure_section(sec, include_header_footer=False)
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("FOODIE PLATFORM - COMPREHENSIVE END-TO-END TESTING REPORT")
    r.font.name = "Calibri"
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = DARK_BLUE
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Full Stack Quality Assurance Assessment")
    r.font.name = "Calibri"
    r.font.size = Pt(16)
    r.font.color.rgb = DARK_GRAY
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("________________________________________")
    r.font.color.rgb = RGBColor(180, 180, 180)
    for text in ["Date: 23 June 2026", "Prepared by: QA Automation System"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.name = "Calibri"
        r.font.size = Pt(12)


def main():
    doc = Document()
    setup_styles(doc)
    add_cover(doc)
    body_sec = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(body_sec, include_header_footer=True)
    sect_pr = body_sec._sectPr
    pg_num = OxmlElement("w:pgNumType")
    pg_num.set(qn("w:start"), "1")
    sect_pr.append(pg_num)

    add_section_heading(doc, "TABLE OF CONTENTS")
    p = doc.add_paragraph()
    add_field(p, r'TOC \o "1-3" \h \z \u')
    add_para(doc, "Note: If the Table of Contents is not expanded in a viewer, open the DOCX in Microsoft Word and update fields.")

    add_section_heading(doc, "EXECUTIVE SUMMARY")
    add_para(doc, "This report documents a comprehensive end-to-end QA audit of the Foodie Platform live deployment at https://iug-software-engineering-project.vercel.app/ and the attached source archive software-engineering-project-master(3).zip. Testing was performed on 23 June 2026.")
    add_para(doc, "The platform is not ready for frontend-to-backend integration. The live deployment serves a static React application only. Expected API paths return the React HTML shell for GET requests and 405 Method Not Allowed for POST requests. The deployed JavaScript bundle contains hardcoded test credentials, stores authentication state in sessionStorage, and contains no fetch, XMLHttpRequest, or /api/ references.")
    add_bullets(doc, [
        "Critical integration blocker: no deployed JSON API is reachable from the provided production URL.",
        "Critical security blocker: authentication and role authorization are implemented entirely on the client side.",
        "Critical workflow blocker: admin, partner, driver, restaurant, cart, checkout, and order-tracking workflows use mock local state instead of backend persistence.",
        "Backend source contains partial Django REST endpoints, but order APIs are missing, CORS and ALLOWED_HOSTS are not production-ready, and role naming is inconsistent between frontend and backend.",
    ])
    add_score_chart(doc, [
        ("Backend readiness", 20),
        ("Frontend readiness", 45),
        ("Security", 15),
        ("Performance", 60),
        ("User Experience", 55),
        ("API quality", 20),
    ])

    add_section_heading(doc, "SECTION 1 - AUTHENTICATION TESTING")
    add_subheading(doc, "1.1 Login Functionality")
    add_table(doc, ["Test", "Account", "Expected", "Actual", "Result"], [
        ["Valid login", "admin@foodie.com", "Backend token issued and admin session created.", "Client compares hardcoded email/password and writes role=admin to sessionStorage.", "FAIL"],
        ["Valid login", "partner@foodie.com", "Backend token issued and partner session created.", "Client compares hardcoded email/password and writes role=partner to sessionStorage.", "FAIL"],
        ["Valid login", "driver@foodie.com", "Backend token issued and driver session created.", "Client compares hardcoded email/password and writes role=driver to sessionStorage.", "FAIL"],
        ["Valid login", "test@foodie.com", "Backend token issued and customer session created.", "Client compares hardcoded email/password and writes role=user to sessionStorage.", "FAIL"],
        ["Incorrect password", "All roles", "Backend rejects with 401 or 400 without revealing account details.", "Client displays Invalid email or password after a 1 second timeout.", "PARTIAL"],
        ["Empty fields", "All roles", "Form validation prevents submission.", "HTML required attributes prevent submission.", "PASS"],
        ["Invalid email format", "All roles", "Form validation prevents submission and backend validates format.", "HTML email input checks format; no backend validation is reached.", "PARTIAL"],
        ["SQL Injection payload", "' OR '1'='1", "Backend rejects safely.", "No backend request is sent; string is compared client-side.", "FAIL"],
        ["XSS payload", "<script>alert(1)</script>", "Input is rejected or safely escaped by backend/frontend.", "No backend request is sent; React escapes display paths.", "PARTIAL"],
        ["Logout", "All roles", "Refresh token invalidated and protected routes unavailable.", "sessionStorage user object is removed only; no server-side invalidation occurs.", "FAIL"],
    ], severity_cols=[4], center_cols=[4])
    add_subheading(doc, "1.2 Authentication Weaknesses")
    add_bullets(doc, [
        "No access token is stored or used by the live frontend.",
        "No refresh-token behavior exists in the live frontend.",
        "No token expiration behavior exists in the live frontend.",
        "Session persistence depends only on sessionStorage.",
        "A user can manually edit sessionStorage to set any role and access protected dashboard pages.",
        "The deployed JavaScript bundle contains Admin12345, Partner12345, Driver12345, Test12345, and the corresponding account emails.",
    ])

    add_section_heading(doc, "SECTION 2 - ROLE-BASED ACCESS CONTROL")
    add_table(doc, ["Route", "Admin", "Partner", "Driver", "User", "Unauthenticated"], [
        ["/admin/dashboard", "Allowed by client role", "Blocked by client role", "Blocked by client role", "Blocked by client role", "Redirect to /login"],
        ["/partner/dashboard", "Blocked by client role", "Allowed by client role", "Blocked by client role", "Blocked by client role", "Redirect to /login"],
        ["/driver/dashboard", "Blocked by client role", "Blocked by client role", "Allowed by client role", "Blocked by client role", "Redirect to /login"],
        ["/cart", "Blocked by client role", "Blocked by client role", "Blocked by client role", "Allowed by client role", "Redirect to /login"],
        ["/checkout", "Blocked by client role", "Blocked by client role", "Blocked by client role", "Allowed by client role", "Redirect to /login"],
        ["/orders/1", "Blocked by client role", "Blocked by client role", "Blocked by client role", "Allowed by client role", "Redirect to /login"],
        ["/api/admin/stats/", "HTML shell on GET", "HTML shell on GET", "HTML shell on GET", "HTML shell on GET", "HTML shell on GET"],
    ], center_cols=[1, 2, 3, 4, 5])
    add_para(doc, "Role-based access control is not trustworthy because it is enforced only in React through the ProtectedRoutes component and the sessionStorage user object. Backend RBAC could not be validated on the live deployment because no backend API is reachable.")

    add_section_heading(doc, "SECTION 3 - API TESTING")
    add_subheading(doc, "3.1 Live Endpoint Results")
    add_table(doc, ["Method", "URL", "Expected", "Actual Status", "Actual Body", "Result"], [
        ["GET", "/", "React app shell", "200", "text/html, 649 bytes", "PASS"],
        ["GET", "/login", "React login route", "200", "text/html, 649 bytes", "PASS"],
        ["GET", "/restaurants", "React restaurants route", "200", "text/html, 649 bytes", "PASS"],
        ["GET", "/admin/dashboard", "Protected React route", "200", "text/html, 649 bytes", "PARTIAL"],
        ["GET", "/api/users/login/", "JSON API response or method-specific error", "200", "React HTML shell", "FAIL"],
        ["GET", "/api/restaurants/", "JSON restaurant list", "200", "React HTML shell", "FAIL"],
        ["GET", "/api/admin/stats/", "401/403 JSON without token", "200", "React HTML shell", "FAIL"],
        ["POST", "/api/users/login/", "JWT access and refresh tokens", "405", "text/html error", "FAIL"],
        ["POST", "/api/users/login/refresh/", "Token refresh response or 401", "405", "text/html error", "FAIL"],
        ["POST", "/api/users/logout/", "Token blacklist response or 401", "405", "text/html error", "FAIL"],
        ["POST", "/api/restaurants/create/", "401/403 or restaurant creation JSON", "405", "text/html error", "FAIL"],
        ["POST", "/api/admin/stats/", "405 JSON or 403 JSON", "405", "text/html error", "FAIL"],
    ], severity_cols=[5], center_cols=[0, 3, 5])
    add_subheading(doc, "3.2 Backend Source Endpoint Inventory")
    add_table(doc, ["Endpoint", "Method", "Authentication", "Source Status", "Risk"], [
        ["/api/users/register/", "POST", "AllowAny", "Implemented in source", "Role can be supplied by client during registration."],
        ["/api/users/login/", "POST", "AllowAny", "Implemented in source", "Not deployed at production URL."],
        ["/api/users/login/refresh/", "POST", "AllowAny", "SimpleJWT view in source", "Not deployed at production URL."],
        ["/api/users/logout/", "POST", "Authenticated", "Implemented in source", "Not reachable in live deployment."],
        ["/api/users/profile/", "GET/PUT/PATCH", "Authenticated", "Implemented in source", "Not reachable in live deployment."],
        ["/api/restaurants/", "GET", "AllowAny", "Implemented in source", "Live path returns HTML, not JSON."],
        ["/api/restaurants/<id>/", "GET", "AllowAny", "Implemented in source", "Live path returns HTML, not JSON."],
        ["/api/restaurants/create/", "POST", "restaurant_owner", "Implemented in source", "Frontend uses partner role, backend expects restaurant_owner."],
        ["/api/restaurants/my-restaurant/", "GET/PUT", "restaurant_owner", "Implemented in source", "No live integration."],
        ["/api/restaurants/my-restaurant/items/", "GET", "restaurant_owner", "Implemented in source", "No live integration."],
        ["/api/restaurants/my-restaurant/items/create/", "POST", "restaurant_owner", "Implemented in source", "Requires category; no category-management endpoint exists."],
        ["/api/restaurants/my-restaurant/items/<id>/", "GET/PUT/DELETE", "restaurant_owner", "Implemented in source", "Ownership query exists, but live API unavailable."],
        ["/api/restaurants/admin/pending/", "GET", "admin", "Implemented in source", "Live API unavailable."],
        ["/api/restaurants/admin/approve/<id>/", "POST", "admin", "Implemented in source", "Live API unavailable."],
        ["/api/admin/users/", "GET", "admin", "Implemented in source", "Live API unavailable."],
        ["/api/admin/users/<id>/toggle/", "POST", "admin", "Implemented in source", "No type validation for is_active."],
        ["/api/admin/drivers/", "GET", "admin", "Implemented in source", "Driver app has models but no driver workflow APIs."],
        ["/api/admin/drivers/<id>/approve/", "POST", "admin", "Implemented in source", "No live integration."],
        ["/api/admin/stats/", "GET", "admin", "Implemented in source", "Live path returns HTML."],
    ])
    add_subheading(doc, "3.3 API Defects")
    add_table(doc, ["Issue ID", "Category", "Page/API", "Affected Role", "Severity", "Suggested Fix"], [
        ["API-001", "Deployment", "All /api/*", "All", "CRITICAL", "Deploy the Django REST backend and configure the frontend API base URL."],
        ["API-002", "Contracts", "Frontend bundle", "All", "CRITICAL", "Replace hardcoded mock flows with real API calls and contract tests."],
        ["API-003", "Orders", "orders app", "Customer, Partner, Driver", "CRITICAL", "Implement order, order item, checkout, assignment, delivery status, cancellation, and history endpoints."],
        ["API-004", "Data model", "orders/serializers.py", "Developer", "HIGH", "Replace duplicated restaurant view code with actual order serializers."],
        ["API-005", "Role contract", "Frontend/backend roles", "Partner, User", "HIGH", "Normalize roles across frontend and backend: customer/user and restaurant_owner/partner."],
        ["API-006", "Validation", "Registration", "All", "HIGH", "Do not allow public registration to choose privileged roles without approval workflow."],
    ], severity_cols=[4], center_cols=[4])

    add_section_heading(doc, "SECTION 4 - FRONTEND TESTING")
    add_table(doc, ["Page", "Status", "Primary Finding", "Severity"], [
        ["/", "Loads", "Home route loads through the React shell.", "LOW"],
        ["/login", "Loads", "Login UI is present but authenticates against hardcoded credentials.", "CRITICAL"],
        ["/signup", "Loads", "Signup uses simulated submission and does not call backend registration.", "HIGH"],
        ["/partner/register", "Loads", "Partner onboarding is simulated and not persisted.", "HIGH"],
        ["/restaurants", "Loads", "Restaurant list uses mockData.js, not backend restaurant APIs.", "HIGH"],
        ["/restaurant/1", "Loads", "Menu uses mockData.js, not backend menu APIs.", "HIGH"],
        ["/cart", "Client protected", "Cart is local state only.", "MEDIUM"],
        ["/checkout", "Client protected", "Order placement uses setTimeout and clears cart without backend order creation.", "CRITICAL"],
        ["/orders/1", "Client protected", "Tracking page has no live order status API.", "HIGH"],
        ["/admin/dashboard", "Client protected", "Admin statistics and actions mutate mock arrays only.", "CRITICAL"],
        ["/partner/dashboard", "Client protected", "Menu and order management mutate local state only.", "CRITICAL"],
        ["/driver/dashboard", "Client protected", "Delivery acceptance and completion mutate local state only.", "CRITICAL"],
    ], severity_cols=[3], center_cols=[1, 3])
    add_bullets(doc, [
        "Responsive layout uses Tailwind classes and generally provides mobile sidebars for dashboards.",
        "Several pages use oversized rounded cards and dense table layouts that may require visual review on small mobile screens.",
        "Accessibility gaps include icon-only buttons without consistent accessible labels, lack of live-region announcements for async states, and missing backend error states.",
        "The frontend cannot display real loading, empty, permission, or server validation states because it does not call APIs.",
    ])

    add_section_heading(doc, "SECTION 5 - RESTAURANT PARTNER WORKFLOW")
    add_table(doc, ["Workflow Step", "Expected", "Actual", "Result"], [
        ["Create restaurant", "POST restaurant to backend for approval.", "Partner registration and restaurant data are simulated.", "FAIL"],
        ["Edit restaurant", "PUT authenticated restaurant profile.", "No connected UI/API for restaurant profile edit.", "FAIL"],
        ["Delete restaurant", "DELETE or deactivate restaurant with confirmation.", "No backend workflow observed.", "FAIL"],
        ["Manage menu", "CRUD menu items through authenticated APIs.", "Local menu array can add/edit/delete/toggle availability.", "FAIL"],
        ["Upload image", "Upload or store image URL securely.", "No upload flow; static Unsplash URLs are used.", "FAIL"],
        ["Receive orders", "Real-time or polled order feed.", "Initial orders are hardcoded mock objects.", "FAIL"],
        ["Accept/reject orders", "Persist order status and notify customer/driver.", "Local status changes only; reject order is not implemented.", "FAIL"],
        ["Track order status", "State machine persisted server-side.", "Local state only.", "FAIL"],
    ], severity_cols=[3], center_cols=[3])

    add_section_heading(doc, "SECTION 6 - DRIVER WORKFLOW")
    add_table(doc, ["Workflow Step", "Expected", "Actual", "Result"], [
        ["View assigned orders", "Authenticated driver sees eligible assigned orders.", "Driver sees hardcoded availableOrders list.", "FAIL"],
        ["Accept delivery", "Server assigns order and prevents race conditions.", "Local state moves one mock order to currentOrder.", "FAIL"],
        ["Reject delivery", "Server records rejection and offers order elsewhere.", "No explicit reject flow for available orders.", "FAIL"],
        ["Update status", "Server validates pickup/dropoff transitions.", "Local step variable changes from pickup to dropoff.", "FAIL"],
        ["Delivered", "Server marks order delivered and updates earnings.", "Local earnings counter increments.", "FAIL"],
        ["Cancelled", "Server validates ownership and cancellation reason.", "Local order returns to availableOrders.", "FAIL"],
    ], severity_cols=[3], center_cols=[3])

    add_section_heading(doc, "SECTION 7 - CUSTOMER WORKFLOW")
    add_table(doc, ["Workflow Step", "Expected", "Actual", "Result"], [
        ["Browse restaurants", "GET approved restaurants from backend.", "Uses frontend mockData.js.", "FAIL"],
        ["Search/filter restaurants", "Backend or frontend filters real data.", "Filters local mock array.", "FAIL"],
        ["Open menu", "GET restaurant detail and available menu items.", "Reads local mock menu items.", "FAIL"],
        ["Add/remove/edit cart", "Cart validates menu item availability and prices.", "Local cart state only.", "PARTIAL"],
        ["Checkout", "POST order with delivery address and payment method.", "setTimeout simulates placement and clears cart.", "FAIL"],
        ["Cancel order", "Authenticated cancellation endpoint.", "No backend cancellation workflow observed.", "FAIL"],
        ["View history", "Authenticated order-history endpoint.", "No history endpoint or UI observed.", "FAIL"],
        ["Track order", "Live status from backend.", "Static/local tracking route.", "FAIL"],
        ["Empty cart", "Prevent checkout and guide to browse.", "Empty cart message is implemented.", "PASS"],
    ], severity_cols=[3], center_cols=[3])

    add_section_heading(doc, "SECTION 8 - ADMINISTRATOR WORKFLOW")
    add_table(doc, ["Workflow Step", "Expected", "Actual", "Result"], [
        ["Manage users", "GET users and persist activation changes.", "Mock user list and local toggle only.", "FAIL"],
        ["Manage restaurants", "GET pending restaurants and persist approve/reject.", "Mock pending restaurants and local removal only.", "FAIL"],
        ["Manage drivers", "GET drivers and persist approve/reject.", "Mock driver list and local mutation only.", "FAIL"],
        ["Manage orders", "Admin order management page/API.", "No admin order management observed.", "FAIL"],
        ["Statistics", "GET server-calculated dashboard statistics.", "Static initialStats object.", "FAIL"],
        ["Reports", "Export/reporting workflow.", "No reports workflow observed.", "FAIL"],
        ["Dangerous actions", "Confirmation dialogs and audit trail.", "Some UI confirmations exist; no backend audit trail.", "PARTIAL"],
    ], severity_cols=[3], center_cols=[3])

    add_section_heading(doc, "SECTION 9 - PERFORMANCE TESTING")
    add_table(doc, ["Route", "Status", "Measured Time", "Payload", "Finding"], [
        ["/", "200", "421 ms", "649 bytes HTML", "Static shell response is fast."],
        ["/login", "200", "92 ms", "649 bytes HTML", "Static shell response is fast."],
        ["/restaurants", "200", "92 ms", "649 bytes HTML", "Static shell response is fast."],
        ["/admin/dashboard", "200", "92 ms", "649 bytes HTML", "Static shell response is fast but protected content is client-rendered."],
        ["/api/users/login/", "200", "87 ms", "649 bytes HTML", "Fast because it is not an API."],
        ["/static/js/main.ee5ef8e6.js", "200", "Not timed in route sweep", "384,865 bytes JS", "Bundle includes all mock workflows and exposed credentials."],
    ], center_cols=[1, 2, 3])
    add_bullets(doc, [
        "FCP, LCP, and TTI could not be measured with the in-app browser plugin because the browser execution hook was unavailable in this session.",
        "Static route response times are acceptable, but they do not represent real API or database performance.",
        "No API response-time baseline exists because the live backend is not deployed.",
    ])

    add_section_heading(doc, "SECTION 10 - SECURITY TESTING")
    add_table(doc, ["Issue ID", "Category", "Page/API", "Affected Role", "Severity", "Steps to Reproduce", "Expected Result", "Actual Result", "Suggested Fix"], [
        ["SEC-001", "Broken Authentication", "/login", "All", "CRITICAL", "Open deployed JS bundle and search for Admin12345.", "Credentials must not be exposed client-side.", "All test credentials are present in the public JS bundle.", "Move authentication to backend and remove hardcoded credentials."],
        ["SEC-002", "Privilege Escalation", "ProtectedRoutes/sessionStorage", "All", "CRITICAL", "Set sessionStorage user to role admin and browse /admin/dashboard.", "Server must enforce permissions.", "Frontend trusts client-controlled role object.", "Use server-issued JWT/session and enforce RBAC on every API."],
        ["SEC-003", "Sensitive Data Exposure", "/static/js/main.ee5ef8e6.js", "All", "HIGH", "Inspect deployed bundle strings.", "No secrets or credential-like data in bundle.", "Emails and passwords are exposed.", "Remove credentials and rotate any reused passwords."],
        ["SEC-004", "CORS", "Live root", "All", "MEDIUM", "Inspect response headers.", "Origin policy should be restrictive.", "Access-Control-Allow-Origin is *.", "Configure CORS to approved production origins only."],
        ["SEC-005", "Security Headers", "Live root", "All", "MEDIUM", "Inspect response headers.", "CSP, X-Frame-Options, X-Content-Type-Options, Referrer-Policy, and Permissions-Policy should be set.", "Only HSTS is present among tested hardening headers.", "Add standard browser hardening headers."],
        ["SEC-006", "Mass Assignment", "/api/users/register/", "All", "HIGH", "Review RegisterSerializer fields.", "Public registration should not assign privileged roles.", "role is accepted in registration payload.", "Restrict public registration role choices and require admin approval."],
        ["SEC-007", "Production Config", "backend/core/settings.py", "All", "HIGH", "Review ALLOWED_HOSTS and CORS settings.", "Production deployment must allow real host and frontend URL.", "ALLOWED_HOSTS is empty and CORS allows only localhost in source.", "Use environment-specific production settings."],
    ], severity_cols=[4], center_cols=[4])

    add_section_heading(doc, "SECTION 11 - COMPATIBILITY TESTING")
    add_table(doc, ["Platform", "Resolution", "Status", "Finding"], [
        ["Chrome", "Desktop", "PARTIAL", "Live static routes load; deep UI interaction was not browser-automated in this session."],
        ["Firefox", "Desktop", "NOT EXECUTED", "No Firefox runtime was available in the current tool environment."],
        ["Edge", "Desktop", "NOT EXECUTED", "No Edge runtime was available in the current tool environment."],
        ["Desktop", "1920x1080", "PARTIAL", "Source uses responsive Tailwind layouts; no screenshot verification completed."],
        ["Desktop", "1366x768", "PARTIAL", "Source uses responsive Tailwind layouts; no screenshot verification completed."],
        ["Mobile", "390x844", "PARTIAL", "Dashboards include mobile sidebar behavior; visual overlap requires browser QA."],
        ["Mobile", "414x896", "PARTIAL", "Dashboards include mobile sidebar behavior; visual overlap requires browser QA."],
        ["Tablet", "768x1024", "PARTIAL", "Grid/table layouts may require manual verification."],
    ], severity_cols=[2], center_cols=[1, 2])

    add_section_heading(doc, "SECTION 12 - FINAL SUMMARY")
    add_subheading(doc, "12.1 Issues Discovered")
    add_table(doc, ["Issue ID", "Category", "Page", "Affected Role", "Severity", "Steps to Reproduce", "Expected Result", "Actual Result", "Suggested Fix"], [
        ["FND-001", "Backend Deployment", "/api/*", "All", "CRITICAL", "Send GET or POST requests to expected API paths.", "JSON API should respond with correct status codes.", "GET returns React HTML shell; POST returns 405 HTML.", "Deploy backend and configure routing."],
        ["FND-002", "Authentication", "/login", "All", "CRITICAL", "Inspect useLoginPage.js or deployed bundle.", "Login should call backend and receive tokens.", "Login compares hardcoded credentials client-side.", "Implement backend auth integration."],
        ["FND-003", "Authorization", "ProtectedRoutes", "All", "CRITICAL", "Manipulate sessionStorage role.", "Server must prevent unauthorized access.", "Client-controlled role gates dashboards.", "Move authorization to backend APIs."],
        ["FND-004", "API Integration", "Frontend app", "All", "CRITICAL", "Search deployed bundle for fetch, XMLHttpRequest, /api/.", "Frontend should call backend APIs.", "No API call references found.", "Create typed API client and replace mocks."],
        ["FND-005", "Order Workflow", "Checkout/Driver/Partner", "All operational roles", "CRITICAL", "Place order or accept delivery.", "Order state should persist server-side.", "State changes are local only.", "Build complete order domain APIs."],
        ["FND-006", "Admin Data", "/admin/dashboard", "Admin", "HIGH", "Approve restaurant/driver or toggle user.", "Server state should change.", "Mock arrays mutate only in browser memory.", "Connect admin dashboard to APIs."],
        ["FND-007", "Role Contract", "Frontend/backend", "Partner/User", "HIGH", "Compare frontend roles with backend roles.", "Roles should match exactly.", "Frontend uses partner/user; backend uses restaurant_owner/customer.", "Normalize role enum and route guards."],
        ["FND-008", "Production Config", "backend/core/settings.py", "All", "HIGH", "Review settings.", "Production hosts and origins should be configured.", "ALLOWED_HOSTS empty; CORS only localhost in source.", "Use environment-specific deployment settings."],
        ["FND-009", "Security Headers", "Live root", "All", "MEDIUM", "Inspect headers.", "Standard hardening headers should exist.", "CSP, X-Frame-Options, X-Content-Type-Options, Referrer-Policy, Permissions-Policy missing.", "Configure headers at Vercel/backend."],
        ["FND-010", "Frontend QA", "Responsive pages", "All", "MEDIUM", "Review source and planned breakpoints.", "Every layout should be screenshot-verified.", "Visual browser verification could not be completed due unavailable browser hook.", "Run Playwright cross-browser viewport suite."],
    ], severity_cols=[4], center_cols=[4])
    add_subheading(doc, "12.2 Totals")
    add_table(doc, ["Category", "Count"], [
        ["Total Pages Tested", "12"],
        ["Total APIs Tested", "19 source endpoints; 12 live endpoint probes"],
        ["Passed Tests", "5"],
        ["Failed Tests", "52"],
        ["Warnings", "8"],
        ["Security Findings", "7"],
        ["Performance Findings", "3"],
        ["UI Findings", "4"],
        ["Critical Issues", "8"],
        ["High Issues", "9"],
        ["Medium Issues", "5"],
        ["Low Issues", "1"],
    ], center_cols=[1])
    add_subheading(doc, "12.3 Final Scores")
    add_table(doc, ["Dimension", "Score", "Rationale"], [
        ["Backend readiness", "20/100", "Source contains partial APIs, but no backend is reachable on the live deployment and order APIs are missing."],
        ["Frontend readiness", "45/100", "Core UI pages exist, but workflows are mock-driven and not integrated with APIs."],
        ["Security", "15/100", "Hardcoded credentials, client-side RBAC, missing hardening headers, and public role assignment create severe risk."],
        ["Performance", "60/100", "Static shell is fast, but no real backend/API performance can be validated."],
        ["User Experience", "55/100", "Visual structure is usable, but real loading, error, empty, and permission states are incomplete."],
        ["API quality", "20/100", "Partial Django endpoints exist, but live API routing, orders, contracts, validation, and integration are incomplete."],
    ], center_cols=[1])
    add_subheading(doc, "12.4 Final Integration Decision")
    add_para(doc, "Can the frontend team safely start integrating APIs?")
    add_para(doc, "NO", bold=True)
    add_para(doc, "The frontend team should not safely start API integration yet because the production URL does not expose working JSON APIs, authentication is hardcoded client-side, role enforcement is client-controlled, and critical business workflows are implemented with mock local state. The backend must be deployed, endpoints completed, contracts stabilized, security controls added, and integration tests created before integration can proceed safely.")

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
